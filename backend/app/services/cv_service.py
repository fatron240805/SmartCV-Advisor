"""CV upload validation, extraction, and section parsing."""

from __future__ import annotations

import re
import os
from dataclasses import dataclass
from datetime import datetime, timezone
from io import BytesIO
from pathlib import Path
from typing import Any
from uuid import uuid4
from zipfile import BadZipFile

from fastapi import HTTPException, UploadFile, status

from app.services.gpt_service import (
    is_gpt_configured,
    parse_cv_image_with_gpt,
    parse_cv_images_with_gpt,
    parse_cv_sections_with_gpt,
)


MAX_CV_SIZE_BYTES = 5 * 1024 * 1024
CONSENT_POLICY_VERSION = "cv-processing-policy-v1"


def get_positive_int_env(name: str, default: int) -> int:
    try:
        return max(1, int(os.getenv(name, str(default))))
    except (TypeError, ValueError):
        return default


def get_positive_float_env(name: str, default: float) -> float:
    try:
        return max(1.0, float(os.getenv(name, str(default))))
    except (TypeError, ValueError):
        return default


GPT_OCR_MAX_PDF_PAGES = get_positive_int_env("GPT_OCR_MAX_PDF_PAGES", 5)
GPT_OCR_PDF_RENDER_ZOOM = get_positive_float_env("GPT_OCR_PDF_RENDER_ZOOM", 2.0)
SUPPORTED_EXTENSIONS = {".pdf", ".doc", ".docx"}
SUPPORTED_IMAGE_EXTENSIONS = {".png", ".jpg", ".jpeg", ".webp", ".bmp"}
SUPPORTED_MIME_TYPES = {
    "application/pdf",
    "application/msword",
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    "image/png",
    "image/x-png",
    "image/jpeg",
    "image/pjpeg",
    "image/webp",
    "image/bmp",
    "image/x-ms-bmp",
    "application/octet-stream",
    "",
}
IMAGE_MIME_TYPES_BY_EXTENSION = {
    ".png": "image/png",
    ".jpg": "image/jpeg",
    ".jpeg": "image/jpeg",
    ".webp": "image/webp",
    ".bmp": "image/bmp",
}

STANDARD_SECTIONS = [
    "Professional Summary",
    "Education",
    "Experience",
    "Projects",
    "Technical Skills",
    "Certifications",
    "Other",
]

SECTION_ALIASES = {
    "Professional Summary": [
        "professional summary",
        "summary",
        "profile",
        "objective",
        "about me",
        "career objective",
        "muc tieu",
        "mục tiêu",
        "gioi thieu",
        "giới thiệu",
        "tom tat",
        "tóm tắt",
    ],
    "Education": [
        "education",
        "academic",
        "hoc van",
        "học vấn",
        "dao tao",
        "đào tạo",
        "university",
    ],
    "Experience": [
        "experience",
        "work history",
        "employment",
        "internship",
        "kinh nghiem",
        "kinh nghiệm",
        "lam viec",
        "làm việc",
    ],
    "Projects": [
        "projects",
        "project",
        "portfolio",
        "academic projects",
        "personal projects",
        "du an",
        "dự án",
        "san pham",
        "sản phẩm",
    ],
    "Technical Skills": [
        "technical skills",
        "skills",
        "technologies",
        "tools",
        "tech stack",
        "ky nang",
        "kỹ năng",
        "cong nghe",
        "công nghệ",
    ],
    "Certifications": [
        "certifications",
        "certificates",
        "licenses",
        "courses",
        "chung chi",
        "chứng chỉ",
        "khoa hoc",
        "khóa học",
    ],
}


def resolve_image_mime_type(extension: str, content_type: str | None) -> str:
    if content_type and content_type.startswith("image/"):
        return "image/jpeg" if content_type == "image/pjpeg" else content_type
    return IMAGE_MIME_TYPES_BY_EXTENSION.get(extension, "image/png")


@dataclass(frozen=True)
class ExtractionResult:
    text: str
    page_count: int | None
    method: str
    language_hints: list[str]
    warnings: list[str]
    quality_score: float
    sections: dict[str, str] | None = None
    section_parser_model: str | None = None
    section_parser_prompt_version: str | None = None


def normalize_text(text: str) -> str:
    cleaned = text.replace("\x00", " ")
    cleaned = re.sub(r"\r\n?", "\n", cleaned)
    cleaned = re.sub(r"[ \t]{2,}", " ", cleaned)
    cleaned = re.sub(r"\n{3,}", "\n\n", cleaned)
    return cleaned.strip()


def normalize_search_text(text: str) -> str:
    lowered = text.lower()
    replacements = {
        "đ": "d",
        "á": "a",
        "à": "a",
        "ả": "a",
        "ã": "a",
        "ạ": "a",
        "ă": "a",
        "ắ": "a",
        "ằ": "a",
        "ẳ": "a",
        "ẵ": "a",
        "ặ": "a",
        "â": "a",
        "ấ": "a",
        "ầ": "a",
        "ẩ": "a",
        "ẫ": "a",
        "ậ": "a",
        "é": "e",
        "è": "e",
        "ẻ": "e",
        "ẽ": "e",
        "ẹ": "e",
        "ê": "e",
        "ế": "e",
        "ề": "e",
        "ể": "e",
        "ễ": "e",
        "ệ": "e",
        "í": "i",
        "ì": "i",
        "ỉ": "i",
        "ĩ": "i",
        "ị": "i",
        "ó": "o",
        "ò": "o",
        "ỏ": "o",
        "õ": "o",
        "ọ": "o",
        "ô": "o",
        "ố": "o",
        "ồ": "o",
        "ổ": "o",
        "ỗ": "o",
        "ộ": "o",
        "ơ": "o",
        "ớ": "o",
        "ờ": "o",
        "ở": "o",
        "ỡ": "o",
        "ợ": "o",
        "ú": "u",
        "ù": "u",
        "ủ": "u",
        "ũ": "u",
        "ụ": "u",
        "ư": "u",
        "ứ": "u",
        "ừ": "u",
        "ử": "u",
        "ữ": "u",
        "ự": "u",
        "ý": "y",
        "ỳ": "y",
        "ỷ": "y",
        "ỹ": "y",
        "ỵ": "y",
    }
    for source, target in replacements.items():
        lowered = lowered.replace(source, target)
    return lowered


def detect_language_hints(text: str) -> list[str]:
    normalized = normalize_search_text(text)
    hints: list[str] = []
    if any(word in normalized for word in ["hoc van", "kinh nghiem", "du an", "ky nang"]):
        hints.append("vi")
    if any(word in normalized for word in ["experience", "education", "projects", "skills"]):
        hints.append("en")
    return hints or ["unknown"]


def validate_file_signature(extension: str, content: bytes) -> None:
    header = content[:8]
    is_valid = False

    if extension == ".pdf":
        is_valid = header.startswith(b"%PDF-")
    elif extension == ".docx":
        is_valid = header.startswith(b"PK\x03\x04")
    elif extension == ".doc":
        is_valid = header.startswith(bytes.fromhex("D0CF11E0A1B11AE1"))
    elif extension in {".png"}:
        is_valid = header.startswith(b"\x89PNG\r\n\x1a\n")
    elif extension in {".jpg", ".jpeg"}:
        is_valid = header.startswith(b"\xff\xd8\xff")
    elif extension == ".webp":
        is_valid = content[:4] == b"RIFF" and content[8:12] == b"WEBP"
    elif extension == ".bmp":
        is_valid = header.startswith(b"BM")

    if not is_valid:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail={
                "code": "CV_FILE_SIGNATURE_INVALID",
                "message": "Tệp không khớp định dạng PDF/DOC/DOCX/ảnh hợp lệ.",
            },
        )


def validate_upload_metadata(file: UploadFile, content: bytes, consent_accepted: bool) -> str:
    if not consent_accepted:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail={
                "code": "CV_CONSENT_REQUIRED",
                "message": "Bạn cần đồng ý xử lý dữ liệu CV để tiếp tục.",
            },
        )

    filename = file.filename or ""
    extension = Path(filename).suffix.lower()
    if extension not in SUPPORTED_EXTENSIONS and extension not in SUPPORTED_IMAGE_EXTENSIONS:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail={
                "code": "CV_FILE_TYPE_NOT_SUPPORTED",
                "message": "Hệ thống chỉ hỗ trợ PDF, DOC, DOCX hoặc ảnh PNG/JPG/JPEG/WEBP/BMP.",
            },
        )

    if file.content_type not in SUPPORTED_MIME_TYPES:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail={
                "code": "CV_MIME_TYPE_INVALID",
                "message": "MIME type của tệp không hợp lệ.",
            },
        )

    if len(content) == 0:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail={"code": "CV_FILE_EMPTY", "message": "Tệp CV đang rỗng."},
        )

    if len(content) > MAX_CV_SIZE_BYTES:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail={
                "code": "CV_FILE_TOO_LARGE",
                "message": "Dung lượng file quá lớn. Giới hạn hiện tại là 5 MB.",
            },
        )

    validate_file_signature(extension, content)
    return extension


def extract_pdf_text(content: bytes, user_id: str) -> ExtractionResult:
    try:
        import fitz
    except ImportError as exc:
        raise HTTPException(
            status_code=status.HTTP_503_SERVICE_UNAVAILABLE,
            detail={
                "code": "CV_PDF_EXTRACTOR_UNAVAILABLE",
                "message": "Máy chủ chưa cài PyMuPDF để đọc PDF.",
            },
        ) from exc

    try:
        document = fitz.open(stream=content, filetype="pdf")
    except Exception as exc:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail={
                "code": "CV_FILE_UNREADABLE",
                "message": "Tệp PDF có thể bị hỏng hoặc không thể đọc.",
            },
        ) from exc

    page_count = document.page_count
    page_texts = [page.get_text("text") for page in document]
    text = normalize_text("\n".join(page_texts))
    warnings: list[str] = []

    if len(text) < 300:
        return extract_pdf_scan_with_gpt(document, user_id=user_id)

    return ExtractionResult(
        text=text,
        page_count=page_count,
        method="pdf_text",
        language_hints=detect_language_hints(text),
        warnings=warnings,
        quality_score=1.0 if len(text) >= 300 else 0.45,
    )


def extract_pdf_scan_with_gpt(document: Any, *, user_id: str) -> ExtractionResult:
    if not is_gpt_configured():
        raise HTTPException(
            status_code=status.HTTP_503_SERVICE_UNAVAILABLE,
            detail={
                "code": "CV_GPT_OCR_UNAVAILABLE",
                "message": "PDF scan cần GPT image model để đọc nội dung. Vui lòng cấu hình OPENAI_API_KEY trong backend/.env.",
            },
        )

    page_count = document.page_count
    page_limit = min(page_count, GPT_OCR_MAX_PDF_PAGES)
    images: list[dict[str, bytes | str]] = []
    try:
        import fitz

        matrix = fitz.Matrix(GPT_OCR_PDF_RENDER_ZOOM, GPT_OCR_PDF_RENDER_ZOOM)
        for page_index in range(page_limit):
            page = document.load_page(page_index)
            pixmap = page.get_pixmap(matrix=matrix, alpha=False)
            images.append(
                {
                    "content": pixmap.tobytes("png"),
                    "mime_type": "image/png",
                    "label": f"Trang {page_index + 1}",
                }
            )
    except Exception as exc:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail={
                "code": "CV_PDF_SCAN_RENDER_FAILED",
                "message": "PDF scan không thể render thành ảnh để GPT đọc nội dung.",
            },
        ) from exc

    gpt_extraction = parse_cv_images_with_gpt(
        images=images,
        user_id=user_id,
        standard_sections=STANDARD_SECTIONS,
    )
    if not gpt_extraction:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail={
                "code": "CV_GPT_OCR_NOT_READABLE",
                "message": "GPT chưa đọc được đủ nội dung từ PDF scan. Vui lòng thử CV rõ nét hơn hoặc PDF có text layer.",
            },
        )

    text = normalize_text(gpt_extraction.raw_text)
    warnings = ["PDF scan được render thành ảnh và đọc trực tiếp bằng GPT image model."]
    if page_count > page_limit:
        warnings.append(f"Hệ thống chỉ gửi {page_limit}/{page_count} trang đầu tiên cho GPT OCR để bảo vệ thời gian xử lý.")
    warnings.extend(gpt_extraction.warnings)
    return ExtractionResult(
        text=text,
        page_count=page_count,
        method="pdf_gpt_ocr",
        language_hints=detect_language_hints(text),
        warnings=warnings,
        quality_score=0.88 if len(text) >= 300 else 0.55,
        sections=gpt_extraction.sections,
        section_parser_model=gpt_extraction.model_version,
        section_parser_prompt_version=gpt_extraction.prompt_version,
    )


def extract_docx_text(content: bytes) -> ExtractionResult:
    try:
        from docx import Document
    except ImportError as exc:
        raise HTTPException(
            status_code=status.HTTP_503_SERVICE_UNAVAILABLE,
            detail={
                "code": "CV_DOCX_EXTRACTOR_UNAVAILABLE",
                "message": "Máy chủ chưa cài python-docx để đọc DOCX.",
            },
        ) from exc

    try:
        document = Document(BytesIO(content))
    except (BadZipFile, ValueError) as exc:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail={
                "code": "CV_FILE_UNREADABLE",
                "message": "Tệp DOCX có thể bị hỏng hoặc không thể đọc.",
            },
        ) from exc

    paragraphs = [paragraph.text for paragraph in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            paragraphs.append(" | ".join(cell.text for cell in row.cells))

    text = normalize_text("\n".join(paragraphs))
    return ExtractionResult(
        text=text,
        page_count=None,
        method="docx",
        language_hints=detect_language_hints(text),
        warnings=[],
        quality_score=1.0 if len(text) >= 300 else 0.65,
    )


def extract_image_text(extension: str, content: bytes, content_type: str | None, user_id: str) -> ExtractionResult:
    if not is_gpt_configured():
        raise HTTPException(
            status_code=status.HTTP_503_SERVICE_UNAVAILABLE,
            detail={
                "code": "CV_GPT_IMAGE_OCR_UNAVAILABLE",
                "message": "Ảnh CV cần GPT image model để đọc nội dung. Vui lòng cấu hình OPENAI_API_KEY trong backend/.env.",
            },
        )

    gpt_extraction = parse_cv_image_with_gpt(
        content=content,
        mime_type=resolve_image_mime_type(extension, content_type),
        user_id=user_id,
        standard_sections=STANDARD_SECTIONS,
    )
    if gpt_extraction:
        warnings = ["Ảnh CV được đọc trực tiếp bằng GPT trước khi chuẩn hóa section."]
        warnings.extend(gpt_extraction.warnings)
        return ExtractionResult(
            text=normalize_text(gpt_extraction.raw_text),
            page_count=1,
            method="image_gpt",
            language_hints=detect_language_hints(gpt_extraction.raw_text),
            warnings=warnings,
            quality_score=0.88 if len(gpt_extraction.raw_text) >= 300 else 0.55,
            sections=gpt_extraction.sections,
            section_parser_model=gpt_extraction.model_version,
            section_parser_prompt_version=gpt_extraction.prompt_version,
        )

    raise HTTPException(
        status_code=status.HTTP_400_BAD_REQUEST,
        detail={
            "code": "CV_GPT_IMAGE_NOT_READABLE",
            "message": "GPT chưa đọc được đủ nội dung từ ảnh CV. Vui lòng thử ảnh rõ nét hơn.",
        },
    )


def extract_cv_text(extension: str, content: bytes, content_type: str | None, user_id: str) -> ExtractionResult:
    if extension == ".pdf":
        extraction = extract_pdf_text(content, user_id=user_id)
    elif extension == ".docx":
        extraction = extract_docx_text(content)
    elif extension in SUPPORTED_IMAGE_EXTENSIONS:
        extraction = extract_image_text(extension, content, content_type, user_id)
    else:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail={
                "code": "CV_LEGACY_DOC_UNSUPPORTED",
                "message": "Định dạng DOC cũ chưa thể đọc ổn định. Vui lòng chuyển CV sang PDF, DOCX hoặc ảnh rõ nét.",
            },
        )

    if len(extraction.text) < 40:
        warning_hint = f" {extraction.warnings[0]}" if extraction.warnings else ""
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail={
                "code": "CV_TEXT_NOT_READABLE",
                "message": f"Không đọc được đủ nội dung CV để phân tích.{warning_hint}",
            },
        )

    return extraction


def classify_heading(line: str) -> str | None:
    normalized = normalize_search_text(line)
    normalized = re.sub(r"[^a-z0-9+#/. ]+", " ", normalized)
    normalized = re.sub(r"\s+", " ", normalized).strip(" :-|")

    if len(normalized) > 60:
        return None

    for section, aliases in SECTION_ALIASES.items():
        if normalized in aliases:
            return section
        if any(normalized.startswith(alias + " ") for alias in aliases):
            return section
    return None


def parse_sections(text: str) -> dict[str, str]:
    sections = {section: "" for section in STANDARD_SECTIONS}
    current_section = "Other"
    buffers = {section: [] for section in STANDARD_SECTIONS}

    for raw_line in text.splitlines():
        line = raw_line.strip()
        if not line:
            if buffers[current_section] and buffers[current_section][-1] != "":
                buffers[current_section].append("")
            continue

        heading = classify_heading(line)
        if heading:
            current_section = heading
            continue

        buffers[current_section].append(line)

    for section, lines in buffers.items():
        sections[section] = normalize_text("\n".join(lines))

    if not any(sections[section] for section in STANDARD_SECTIONS if section != "Other"):
        sections["Other"] = text

    return sections


async def build_cv_document(
    *,
    file: UploadFile,
    content: bytes,
    consent_accepted: bool,
    user_id: str,
    policy_version: str | None,
) -> dict[str, Any]:
    extension = validate_upload_metadata(file, content, consent_accepted)
    extraction = extract_cv_text(extension, content, file.content_type, user_id)

    if extraction.sections and any(value.strip() for value in extraction.sections.values()):
        sections = extraction.sections
        section_parser = "gpt_image" if extraction.method == "image_gpt" else "gpt_ocr"
        section_parser_model = extraction.section_parser_model
        section_parser_prompt_version = extraction.section_parser_prompt_version
    else:
        gpt_sections = parse_cv_sections_with_gpt(
            raw_text=extraction.text,
            user_id=user_id,
            standard_sections=STANDARD_SECTIONS,
        )
        if gpt_sections:
            sections = gpt_sections.sections
            section_parser = "gpt"
            section_parser_model = gpt_sections.model_version
            section_parser_prompt_version = gpt_sections.prompt_version
        else:
            sections = parse_sections(extraction.text)
            section_parser = "rule_based"
            section_parser_model = None
            section_parser_prompt_version = None
    now = datetime.now(timezone.utc)
    cv_id = f"CV_{uuid4().hex[:10].upper()}"
    filename = file.filename or f"uploaded-cv{extension}"

    return {
        "_id": cv_id,
        "Loai": extension.lstrip("."),
        "DungLuong": len(content),
        "TenFileGoc": filename,
        "TenFileNormalized": normalize_search_text(filename),
        "ContentType": file.content_type,
        "NgayTaiLen": now,
        "TrangThai": "uploaded",
        "MaNganh": None,
        "MaKH": user_id,
        "Consent": {
            "accepted": True,
            "policy_version": policy_version or CONSENT_POLICY_VERSION,
            "accepted_at": now,
            "user_id": user_id,
        },
        "Extraction": {
            "text": extraction.text,
            "text_length": len(extraction.text),
            "sections": sections,
            "page_count": extraction.page_count,
            "method": extraction.method,
            "language_hints": extraction.language_hints,
            "warnings": extraction.warnings,
            "quality_score": extraction.quality_score,
            "section_parser": section_parser,
            "section_parser_model": section_parser_model,
            "section_parser_prompt_version": section_parser_prompt_version,
        },
    }


def format_file_size(size_bytes: int) -> str:
    if size_bytes >= 1024 * 1024:
        return f"{size_bytes / (1024 * 1024):.1f} MB"
    return f"{max(1, round(size_bytes / 1024))} KB"


def public_cv_metadata(cv: dict[str, Any]) -> dict[str, Any]:
    extraction = cv.get("Extraction", {})
    return {
        "cv_id": cv.get("_id"),
        "filename": cv.get("TenFileGoc"),
        "file_type": cv.get("Loai"),
        "size_bytes": cv.get("DungLuong", 0),
        "size_label": format_file_size(int(cv.get("DungLuong", 0) or 0)),
        "status": cv.get("TrangThai"),
        "uploaded_at": cv.get("NgayTaiLen"),
        "target_role_id": cv.get("MaNganh"),
        "consent": cv.get("Consent", {}),
        "extraction": {
            "text_length": extraction.get("text_length", 0),
            "page_count": extraction.get("page_count"),
            "method": extraction.get("method"),
            "language_hints": extraction.get("language_hints", []),
            "warnings": extraction.get("warnings", []),
            "quality_score": extraction.get("quality_score", 0),
            "section_parser": extraction.get("section_parser", "rule_based"),
            "section_parser_model": extraction.get("section_parser_model"),
            "section_parser_prompt_version": extraction.get("section_parser_prompt_version"),
        },
    }
